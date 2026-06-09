from docx import Document

def fix_template():
    try:
        doc = Document('C:/Users/Arunprakash/OneDrive/Documents/FACULTY_TRACK.docx')
        
        # Replace specific identifiable fields with [[tags]]
        for p in doc.paragraphs:
            if "FACULTY TRACK" in p.text:
                p.text = "[[title]]"
            elif "ARUNPRAKASH" in p.text:
                p.text = "[[author_names]]"
            elif len(p.text) > 150:
                p.text = "[[body]]"
                
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "FACULTY TRACK" in cell.text:
                        cell.text = "[[title]]"
                    elif "ARUNPRAKASH" in cell.text:
                        cell.text = "[[author_names]]"
                        
        doc.save('public/FACULTY_TRACK_TAGGED.docx')
        print("Successfully created public/FACULTY_TRACK_TAGGED.docx")
    except Exception as e:
        print("Error:", e)

if __name__ == '__main__':
    fix_template()
